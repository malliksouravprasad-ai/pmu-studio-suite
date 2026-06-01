"""Word report generator for APP-008."""
import io
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .deliverable_model import ReportConfig, SectionDef

_NAVY  = RGBColor(0x1F, 0x38, 0x64)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT = RGBColor(0xD9, 0xE1, 0xF2)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_table(doc: Document, df: pd.DataFrame):
    if df.empty:
        doc.add_paragraph("(No data available for this section.)")
        return

    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for ci, col_name in enumerate(df.columns):
        cell = hdr.cells[ci]
        _set_cell_bg(cell, "1F3864")
        run = cell.paragraphs[0].add_run(str(col_name))
        run.font.bold = True
        run.font.color.rgb = _WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, (_, row) in enumerate(df.iterrows(), start=1):
        is_total = str(row.iloc[0]) == "TOTAL" if len(row) > 0 else False
        for ci, (col, val) in enumerate(row.items()):
            cell = table.rows[ri].cells[ci]
            text = str(val) if pd.notna(val) else ""
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(9)
            if is_total:
                _set_cell_bg(cell, "BDD7EE")
                run.font.bold = True
            elif ri % 2 == 0:
                _set_cell_bg(cell, "D9E1F2")

    doc.add_paragraph()  # spacing after table


def save_word_report(
    section_data: dict, config: ReportConfig, artifact_id: str
) -> tuple:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config.report_title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = _NAVY

    for label, value in [
        ("Programme", config.programme),
        ("Organisation", config.organization),
        ("Prepared by", config.author),
        ("Date", config.report_date),
    ]:
        if value:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_lbl = p2.add_run(f"{label}: ")
            r_lbl.font.bold = True
            r_lbl.font.size = Pt(11)
            r_val = p2.add_run(value)
            r_val.font.size = Pt(11)

    if config.description:
        doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.add_run(config.description).font.size = Pt(10)

    doc.add_paragraph()
    p_art = doc.add_paragraph()
    art_run = p_art.add_run(f"Artifact ID: {artifact_id}")
    art_run.font.italic = True
    art_run.font.size = Pt(9)

    doc.add_page_break()

    # Sections
    for sdef in config.sections:
        if not sdef.enabled:
            continue

        doc.add_heading(sdef.title, level=1)

        if sdef.section_type == "narrative":
            if sdef.narrative:
                p = doc.add_paragraph(sdef.narrative)
                p.runs[0].font.size = Pt(11)

        elif sdef.section_type in ("table", "highlights"):
            if sdef.narrative:
                p = doc.add_paragraph(sdef.narrative)
                if p.runs:
                    p.runs[0].font.size = Pt(11)
            df = section_data.get(sdef.section_id, pd.DataFrame())
            _add_table(doc, df)

        doc.add_paragraph()  # spacing between sections

    fname = f"{artifact_id}_report.docx"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return fname, buf
