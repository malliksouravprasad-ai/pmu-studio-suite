"""
Document Generator — converts markdown guides to styled DOCX files.

Run from PMU_Tools root:
    .venv\\Scripts\\python docs\\generate_docs.py
"""
import sys, os, re
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from shared import register, generate_id, timestamp_suffix

PMU_ROOT = Path(__file__).parent.parent

# ── Brand colors ──────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE = RGBColor(0xD6, 0xDC, 0xE4)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
GREY_TEXT  = RGBColor(0x60, 0x60, 0x60)


# ── Document setup ────────────────────────────────────────────────────────────

def _new_doc() -> Document:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    return doc


def _set_para_spacing(para, before: int = 0, after: int = 6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def _shade_cell(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


# ── Content writers ───────────────────────────────────────────────────────────

def _add_cover(doc: Document, title: str, subtitle: str, version: str):
    doc.add_paragraph()  # top space
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OSEPA PMU Tool Suite")
    run.font.color.rgb = MED_BLUE
    run.font.size = Pt(12)
    run.font.bold = True
    _set_para_spacing(p, 60, 4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.color.rgb = DARK_BLUE
    run2.font.size = Pt(22)
    run2.font.bold = True
    _set_para_spacing(p2, 4, 8)

    if subtitle:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(subtitle)
        run3.font.color.rgb = GREY_TEXT
        run3.font.size = Pt(12)
        _set_para_spacing(p3, 0, 60)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = meta.add_run(f"Version {version}  ·  OSEPA PMU Tool Suite")
    run4.font.color.rgb = GREY_TEXT
    run4.font.size = Pt(9)
    _set_para_spacing(meta, 0, 4)

    doc.add_page_break()


def _add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    _set_para_spacing(p, 14, 6)
    # Horizontal rule effect via bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = MED_BLUE
    _set_para_spacing(p, 10, 3)
    return p


def _add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    _set_para_spacing(p, 6, 2)
    return p


def _add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    _set_para_spacing(p, 0, 4)
    return p


def _add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    # Parse bold (**text**)
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(10.5)
        if i % 2 == 1:
            run.font.bold = True
    p.paragraph_format.left_indent = Cm(level * 0.5)
    _set_para_spacing(p, 0, 2)
    return p


def _add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.left_indent = Cm(0.8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    _set_para_spacing(p, 1, 1)
    return p


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _shade_cell(cell, "1F3864")
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            _shade_cell(cell, fill)
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(10)

    doc.add_paragraph()  # space after table


def _add_note(doc: Document, text: str, note_type: str = "info"):
    colors = {"info": "D6DCE4", "warning": "FFF2CC", "tip": "E2EFDA"}
    fill = colors.get(note_type, "D6DCE4")
    labels = {"info": "NOTE", "warning": "IMPORTANT", "tip": "TIP"}
    label = labels.get(note_type, "NOTE")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_label = p.add_run(f"{label}: ")
    run_label.font.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = DARK_BLUE
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    p._p.get_or_add_pPr().append(shading)
    _set_para_spacing(p, 4, 4)


# ── Setup Guide ───────────────────────────────────────────────────────────────

def build_setup_guide() -> Document:
    doc = _new_doc()
    _add_cover(doc, "PMU Tool Suite", "Setup Guide", "1.0")

    _add_h1(doc, "1.  What Is the PMU Tool Suite?")
    _add_body(doc, "The PMU Tool Suite is a collection of standalone productivity applications built for Project Monitoring Unit (PMU) staff. Each application takes a data file as input, processes it, and produces a ready-to-use output — reports, dashboards, forms, presentations, or registers.")
    _add_body(doc, "No database. No server. No login system. Each app runs locally on your computer and produces files you can immediately open in Excel, Word, PowerPoint, or Google Sheets.")

    _add_h1(doc, "2.  System Requirements")
    _add_table(doc,
        ["Requirement", "Minimum"],
        [
            ["Operating System", "Windows 10 or Windows 11"],
            ["Python", "Version 3.10 or higher (3.14 recommended)"],
            ["RAM", "4 GB"],
            ["Disk Space", "2 GB free"],
            ["Internet", "Required only for Google integration features"],
            ["Browser", "Chrome, Edge, or Firefox (for Streamlit apps)"],
        ]
    )

    _add_h1(doc, "3.  Folder Structure")
    for line in [
        "PMU_Tools/",
        "  .venv/              Python environment (do not modify)",
        "  credentials/        Google login credentials (set up once)",
        "  shared/             Common services used by all apps",
        "  apps/",
        "    APP-001_Form_Builder/",
        "    APP-002_Data_Cleaner/",
        "    ... (more apps)",
        "  templates/          Shared Excel, Word, PPT templates",
        "  inputs/             Drop your raw data files here",
        "  outputs/            All generated files appear here",
        "  registry.csv        Automatic log of every output produced",
    ]:
        _add_code(doc, line)

    _add_h1(doc, "4.  First-Time Installation")
    _add_h2(doc, "Step 1 — Open the PMU Tools folder")
    _add_body(doc, "Navigate to the folder where PMU Tools is saved on your computer.")

    _add_h2(doc, "Step 2 — Check that the virtual environment exists")
    _add_body(doc, "Look for a folder named .venv inside PMU_Tools. If it exists, setup is done — skip to Step 4.")

    _add_h2(doc, "Step 3 — Create the environment (if .venv is missing)")
    _add_body(doc, "Open Command Prompt in the PMU_Tools folder and run:")
    _add_code(doc, "python -m venv .venv")
    _add_code(doc, ".venv\\Scripts\\pip install -r requirements.txt")
    _add_body(doc, "This installs all required packages. It takes 2–3 minutes on first run.")

    _add_h2(doc, "Step 4 — Verify installation")
    _add_code(doc, ".venv\\Scripts\\python verify_shared.py")
    _add_body(doc, "You should see: === ALL 5 COMMON SERVICES VERIFIED ===")

    _add_h1(doc, "5.  Google Integration Setup (Optional)")
    _add_body(doc, "Google integration allows apps to read data from Google Sheets and publish forms, reports, or datasets directly to Google Drive.")
    _add_note(doc, "You only need to do this once. All 14 apps share the same credentials.", "tip")

    _add_h2(doc, "Step 1 — Create a Google Cloud Project")
    for item in [
        "Go to: https://console.cloud.google.com",
        "Sign in with the Google account that will own your forms and sheets",
        "Click New Project → Name it PMU Tools → Click Create",
    ]:
        _add_bullet(doc, item)

    _add_h2(doc, "Step 2 — Enable APIs")
    _add_body(doc, "In your new project, go to APIs & Services → Library and enable:")
    for api in ["Google Forms API", "Google Sheets API", "Google Drive API"]:
        _add_bullet(doc, api)

    _add_h2(doc, "Step 3 — Create Credentials")
    for item in [
        "Go to APIs & Services → Credentials",
        "Click Create Credentials → OAuth 2.0 Client ID",
        "Application type: Desktop App  |  Name: PMU Tools",
        "Click Create → Download JSON",
    ]:
        _add_bullet(doc, item)

    _add_h2(doc, "Step 4 — Place the credentials file")
    for item in [
        "Rename the downloaded file to: credentials.json",
        "Copy it to: PMU_Tools/credentials/credentials.json",
    ]:
        _add_bullet(doc, item)

    _add_h2(doc, "Step 5 — First login")
    _add_body(doc, "The first time any app uses a Google feature, a browser window will open. Sign in with the same Google account used in Step 1. A token.json file is saved automatically — you will not be asked again.")
    _add_note(doc, "Do not share credentials.json or token.json with others. These files give full access to your Google account.", "warning")

    _add_h1(doc, "6.  Launching an App")
    _add_body(doc, "Each app has a run.bat file. Double-click it to launch the app in your browser at http://localhost:8501")
    _add_body(doc, "To stop the app, close the black Command Prompt window.")

    _add_h1(doc, "7.  Understanding Outputs")
    _add_body(doc, "All outputs are saved in: PMU_Tools/outputs/<app-folder>/")
    _add_body(doc, "Every output has a unique Artifact ID in the format:")
    _add_code(doc, "PMU-2026-FLN-SURVEY-00001")
    _add_code(doc, "PMU-2026-ADM-REPORT-00005")

    _add_h1(doc, "8.  The Registry (registry.csv)")
    _add_body(doc, "Every time an app generates an output, it records the event in registry.csv. Open this file in Excel to see all outputs produced, filter by project or app, and track review status.")
    _add_table(doc,
        ["Column", "Description"],
        [
            ["app_id", "Which app produced the output (e.g. APP-001)"],
            ["artifact_id", "Unique ID of the output"],
            ["date_generated", "Date and time"],
            ["project", "Project name or code"],
            ["report_type", "Type of output (Form, Report, Dictionary, etc.)"],
            ["output_file", "Full path to the generated file"],
            ["status", "Generated / Reviewed / Submitted / Archived"],
        ]
    )

    _add_h1(doc, "9.  Troubleshooting")
    _add_table(doc,
        ["Problem", "Solution"],
        [
            ["App does not open", "Check that .venv exists. Re-run pip install -r requirements.txt"],
            ["Module not found error", "Run Python from .venv\\Scripts\\python.exe"],
            ["Google login fails", "Delete credentials/token.json and try again"],
            ["Output file not saved", "Check that outputs/ folder exists and is not read-only"],
            ["Registry not updating", "Ensure registry.csv is not open in Excel while app is running"],
            ["Browser does not open", "Manually go to http://localhost:8501"],
        ]
    )

    return doc


# ── APP-001 User Guide ────────────────────────────────────────────────────────

def build_app001_guide() -> Document:
    doc = _new_doc()
    _add_cover(doc, "APP-001 Form Builder", "User Guide", "1.0")

    _add_h1(doc, "1.  What This App Does")
    _add_body(doc, "The Form Builder lets PMU staff design monitoring forms without starting from scratch every time. You use a visual interface to add questions, arrange them into sections, set validation rules, and generate a live Google Form and a Data Dictionary in one click.")
    _add_table(doc,
        ["Before this tool", "After this tool"],
        [
            ["Each form designed manually in Google Forms", "Forms built from reusable templates"],
            ["No standardization across teams", "Consistent structure for all forms"],
            ["No documentation of form fields", "Automatic Data Dictionary produced with every form"],
            ["Questions recreated from scratch each time", "Question bank stores 20+ common questions ready to use"],
        ]
    )

    _add_h1(doc, "2.  When to Use This App")
    for item in [
        "Create a new monitoring, survey, or observation form",
        "Adapt an existing form for a new program or district",
        "Standardize a form across multiple teams",
        "Document what every form field means (for data quality review)",
        "Create a Google Form that automatically collects responses into a linked sheet",
    ]:
        _add_bullet(doc, item)

    _add_h1(doc, "3.  What You Get as Outputs")
    _add_table(doc,
        ["Output", "Format", "Description"],
        [
            ["Google Form", "Live link", "Ready for respondents — shareable on WhatsApp"],
            ["Response Sheet", "Google Sheet", "Auto-linked — responses appear here in real time"],
            ["Data Dictionary", "Excel (.xlsx)", "Documents every question, type, options, and rules"],
        ]
    )

    _add_h1(doc, "4.  Launching the App")
    for step in [
        "Open PMU_Tools/apps/APP-001_Form_Builder/",
        "Double-click run.bat",
        "The app opens in your browser at http://localhost:8501",
    ]:
        _add_bullet(doc, step)
    _add_note(doc, "If the browser does not open automatically, go to http://localhost:8501 manually. To stop the app, close the black Command Prompt window.", "info")

    _add_h1(doc, "5.  App Navigation")
    _add_body(doc, "The app has 4 pages, accessed from the sidebar on the left:")
    _add_table(doc,
        ["Page", "Purpose"],
        [
            ["Home", "Dashboard: create new form or load a template"],
            ["Form Designer", "Build and edit your form — the main workspace"],
            ["Template Library", "Browse, load, and save form templates"],
            ["Question Bank", "Reusable questions — add them to any form"],
            ["Generate Output", "Export to Google Form and Data Dictionary"],
        ]
    )

    _add_h1(doc, "6.  Form Designer — Main Builder")

    _add_h2(doc, "Top Bar")
    _add_body(doc, "Fill in the Form Title, Project Code, and optional Description. Use 💾 Save to save the form as a template, or 🗑 Clear to start over.")

    _add_h2(doc, "Left Panel — Adding Questions")
    _add_table(doc,
        ["Button", "Question Type", "Use When"],
        [
            ["T  Short Text", "Short Text", "Name, code, district, block"],
            ["P  Paragraph", "Long Text", "Observations, remarks, comments"],
            ["◉  Single Select", "Dropdown", "Yes/No, ratings, categories"],
            ["☑  Multi Select", "Checkboxes", "Select all that apply"],
            ["#  Numeric", "Number entry", "Counts, percentages, enrolment"],
            ["📅 Date", "Date picker", "Visit date, reporting date"],
            ["⊞  Matrix Grid", "Rating grid", "Rate multiple items on one scale"],
        ]
    )
    _add_body(doc, "Bulk Import: Paste a list of question labels (one per line) and click Import — all are added as Short Text questions in one step.")

    _add_h2(doc, "Center Panel — Form Canvas")
    for item in [
        "Click any question to select it and edit its properties on the right",
        "Use ↑ ↓ arrows to reorder questions within a section",
        "Use ✕ to delete a question",
        "Use ➕ Add to Section Name to add directly under a section",
    ]:
        _add_bullet(doc, item)

    _add_h2(doc, "Right Panel — Question Properties")
    _add_table(doc,
        ["Property", "Description"],
        [
            ["Label", "The question text shown to respondents (required)"],
            ["Type", "Change the question type anytime"],
            ["Help Text", "Guidance shown below the question"],
            ["Required", "Tick to make the question mandatory (marked ✱)"],
            ["Options", "For Single/Multi Select: enter one option per line"],
            ["Min / Max", "For Numeric: set the allowed value range"],
            ["Rows / Columns", "For Matrix: define items to rate and rating labels"],
            ["Show/Hide Logic", "Show this question only if another equals a specific value"],
        ]
    )

    _add_h1(doc, "7.  Built-in Templates")
    _add_table(doc,
        ["Template", "Sections", "Questions", "Best For"],
        [
            ["School Visit Monitoring Form", "4", "20", "Field visit monitoring"],
            ["Training Attendance & Feedback", "2", "11", "Training programs"],
            ["FLN Classroom Observation Form", "4", "15", "FLN classroom visits"],
        ]
    )
    _add_body(doc, "To use a template: go to Template Library → click Load into Form Designer → edit as needed.")

    _add_h1(doc, "8.  Generate Output")
    _add_h2(doc, "Step 1 — Check the summary")
    _add_body(doc, "The Generate Output page shows the number of sections, questions, and required fields. Review before generating.")

    _add_h2(doc, "Step 2A — Data Dictionary")
    _add_body(doc, "Click Generate Data Dictionary (XLSX). A formatted Excel file is produced immediately documenting every question. No internet required.")

    _add_h2(doc, "Step 2B — Google Form")
    _add_note(doc, "Requires credentials.json in PMU_Tools/credentials/. See Setup Guide Section 5 for one-time Google setup.", "warning")
    _add_body(doc, "Click 🚀 Create Google Form. The app creates a live Google Form with all sections, questions, and validation, and links a Google Sheet for response collection. Two links are returned: the Form link (share with respondents) and the Response Sheet link.")

    _add_h1(doc, "9.  Step-by-Step: Build Your First Form (10 Minutes)")
    steps = [
        ("Open the app", "Double-click run.bat → app opens in browser"),
        ("Create a new form", "Click ➕ Create New Form on the Home page"),
        ("Set the title", "In Form Designer: type the form name and project code in the top bar"),
        ("Add a section", "Sections are pre-created. Add more with ➕ Add Section if needed"),
        ("Add questions", "Click question types in the left panel — each click adds one question"),
        ("Edit each question", "Click the question in the canvas → edit label, type, options in the right panel"),
        ("Mark required fields", "Tick the Required checkbox for all mandatory questions"),
        ("Save the form", "Click 💾 Save to add it to the Template Library"),
        ("Generate outputs", "Go to Generate Output → click Generate Data Dictionary → download"),
        ("Create Google Form", "Click 🚀 Create Google Form → copy the link and share on WhatsApp"),
    ]
    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run1 = p.add_run(f"  {i}.  {title} — ")
        run1.font.bold = True
        run1.font.size = Pt(10.5)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10.5)
        _set_para_spacing(p, 0, 3)

    _add_h1(doc, "10.  Frequently Asked Questions")
    faqs = [
        ("Can I have multiple sections?",
         "Yes. Click ➕ Add Section. Each section becomes a separate page in Google Forms."),
        ("What happens if I close the app?",
         "Your form is saved only when you click 💾 Save. If you did not save before closing, the form is lost. Always save before closing."),
        ("Can multiple people build forms at the same time?",
         "No. The app runs on one computer at a time. Share completed forms as templates or as Google Form links."),
        ("The Google Form is missing some questions — why?",
         "If the form has 50+ questions, some may be skipped by the API. Split into smaller forms or contact the PMU tool team."),
        ("Can I edit a Google Form after it is created?",
         "Yes — open it in Google Forms and edit directly. Regenerate the Data Dictionary from the app if the structure changes significantly."),
        ("What is the Artifact ID?",
         "Every output gets a unique ID like PMU-2026-FLN-FORM-00001, which is recorded in registry.csv for tracking."),
    ]
    for q_text, a_text in faqs:
        p = doc.add_paragraph()
        run_q = p.add_run(f"Q: {q_text}")
        run_q.font.bold = True
        run_q.font.size = Pt(10.5)
        run_q.font.color.rgb = DARK_BLUE
        _set_para_spacing(p, 6, 1)
        p2 = doc.add_paragraph()
        run_a = p2.add_run(f"A: {a_text}")
        run_a.font.size = Pt(10.5)
        _set_para_spacing(p2, 0, 6)

    _add_h1(doc, "11.  Output Files Reference")
    _add_table(doc,
        ["File", "Location", "When Produced"],
        [
            ["Data Dictionary (.xlsx)", "outputs/APP-001_Form_Builder/", "Clicking Generate Data Dictionary"],
            ["Saved templates (.json)", "apps/APP-001_Form_Builder/data/templates/", "Clicking 💾 Save in Form Designer"],
            ["Registry entry", "registry.csv (root)", "Automatically on every output"],
        ]
    )

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    out_dir = PMU_ROOT / "outputs" / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)

    ts = timestamp_suffix()

    print("Building Setup Guide...")
    setup_doc = build_setup_guide()
    setup_path = out_dir / f"PMU_Tool_Suite_Setup_Guide_{ts}.docx"
    setup_doc.save(str(setup_path))
    aid1 = generate_id("DOC", "SETUP")
    register("DOCS", aid1, "INTERNAL", "Setup Guide", str(setup_path))
    print(f"  Saved: {setup_path}")

    print("Building APP-001 User Guide...")
    app001_doc = build_app001_guide()
    app001_path = out_dir / f"APP-001_Form_Builder_User_Guide_{ts}.docx"
    app001_doc.save(str(app001_path))
    aid2 = generate_id("DOC", "GUIDE")
    register("DOCS", aid2, "INTERNAL", "User Guide - APP-001", str(app001_path))
    print(f"  Saved: {app001_path}")

    print(f"\nDone. Both documents saved to: {out_dir}")
    print(f"  {aid1} — Setup Guide")
    print(f"  {aid2} — APP-001 User Guide")


if __name__ == "__main__":
    main()
